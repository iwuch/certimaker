from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import pinyin
import datetime

BLACK = RGBColor(0,0,0)
RED = RGBColor(255,0,0)
GREEN = RGBColor(0,255,0)
BLUE = RGBColor(0,0,255)
TODATE = datetime.datetime.now().strftime('%Y.%m.%d') 
THISY = datetime.datetime.now().strftime('%Y')
LASTY = str(int(THISY)-1)
class DocWriter:
	def __init__(self,name,award):
		self.doc = Document()
		self.name = name
		self.award = award
	def new_para(self):
		self.para = self.doc.add_paragraph()
	def write_run(self, content, fontsize = 25, fontname = '华文中宋',
		alignment = WD_ALIGN_PARAGRAPH.LEFT, color = BLACK, 
		underline = False, bold = False):
		para = self.para
		run = para.add_run(content)
		run.font.size = Pt(fontsize)
		run.font.name = fontname
		run.font.underline = underline
		run.font.bold = bold
		r = run._element.rPr.rFonts
		r.set(qn('w:eastAsia'),fontname)
		para.alignment = alignment
		run.font.color.rgb = color
	def signature(self,sign):
		self.new_para()
		self.write_run(sign,alignment=WD_ALIGN_PARAGRAPH.RIGHT)
		self.new_para()
		self.write_run(TODATE,alignment=WD_ALIGN_PARAGRAPH.RIGHT)
	def save_doc(self):
		pyname = pinyin.get(self.name, format='strip', delimiter="")
		pyaward= pinyin.get(self.award,format='strip', delimiter="")
		filename = pyname.upper() + '_'+ pyaward + '.docx'
		self.doc.save(filename)

def make_certificate(name,award):
	awardoc = DocWriter(name,award)
	awardoc.new_para()
	awardoc.write_run('奖  状', fontsize = 50,alignment= WD_ALIGN_PARAGRAPH.CENTER,color = RED,bold=True)
	awardoc.new_para()
	awardoc.write_run(' '*2+name+' '*2,fontsize = 30,underline = True,bold=True)
	awardoc.write_run(' 同学：',fontsize = 30)
	awardoc.new_para()
	awardoc.write_run(f'\t你在{LASTY}—{THISY}年度表现优异，被授予')
	awardoc.write_run(' '*2 + award + ' '*2, underline=True, bold=True)
	awardoc.write_run('荣誉称号。')
	awardoc.new_para()
	awardoc.write_run('\t特发此证，以资鼓励！')
	awardoc.signature('python大学')
	awardoc.save_doc()

if __name__ == '__main__':
	make_certificate('王大雷','三好学生')